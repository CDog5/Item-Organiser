import pandas as pd
from docx import Document
import os
from docx.shared import Inches
def create_xlsx_data():
    with open("qrdata.txt","r") as f:
        lines = f.readlines()
    qrs=[]
    items=[]
    #split the data into the columns
    for line in lines:
        x,y = line.split(",")
        qrs.append(int(x))
        items.append(y)


    df = pd.DataFrame({'QR ID': qrs,
                       'Item': items})

    # Write data to excel
    df.to_excel("Data.xlsx", sheet_name="QR DATA", index=False)
def create_qr_docx(size,fr,to):
    #fdr is the path to the folder where you want to generate your qr codes
    fdr = ""
    document = Document()
    p = document.add_paragraph()
    r = p.add_run()
    a = []
    for i in range(fr,to+1):
        for filename in os.listdir(fdr):
            if filename.endswith(".png"):
              no = filename.replace(".png","")
              if i == int(no):
                    r.add_picture(os.path.join(fdr,filename),height=Inches(size),width=Inches(size))
                    r.add_text(no)

    document.save("QRC.docx")
